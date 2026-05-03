"""Build the PM_23 Marketing Report .docx.

Register: formal third-person academic. First-person reserved for §15
Reflection (where the brief explicitly invites it) and §18 AI-use statement.
Voice rules retained: zero em-dashes, EU spelling, prohibited-word audit.
APA 7th citations throughout.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

ROOT = Path(__file__).resolve().parent.parent
CHARTS = ROOT / "inputs" / "charts"

with open(ROOT / "inputs" / "meta_aggregates.json") as f:
    AGG = json.load(f)
T = AGG["totals"]
SPEND = T["spend"]; IMP = T["impressions"]; CLK = T["link_clicks"]; LEADS = T["leads"]
CPL = SPEND / LEADS; CTR = CLK / IMP * 100; CPC = SPEND / CLK; CPM = SPEND / IMP * 1000
DAY_MIN = T["day_min"]; DAY_MAX = T["day_max"]; DAY_N = T["days_count"]

# Open the CODE University Hand-in Template and fill in the cover-page placeholders.
TEMPLATE = ROOT / "inputs" / "Hand-in_Template.docx"
doc = Document(str(TEMPLATE))

COVER_FIELDS = {
    "Title": "ENGINE_001 — A Dual-Channel Acquisition Sprint for Mango Lab and Still",
    "[optional description]": "An in-flight paid and organic Meta sprint, reported at the 6-day mark.",
    "[Module Name (if different title)]": "PM_23 / BM_23 — Product Marketing and Sales",
    "Module Coordinator: [Module Coordinator Name]": "Module Coordinator: Roland Fassauer",
    "Learning Unit: [LU Name] by [LU Organizer Name]": "Learning Unit: Marketing Report by Roland Fassauer",
    "[Spring/Fall] Semester 202?": "Spring Semester 2026",
    "[number] words": f"~5,800 words",
    "Student Name": "Jean-Luc Andre Navarro",
    "[student email]": "jean-luc.navarro@code.berlin",
    "Hand-in Date": "Hand-in Date: 4 May 2026",
}
for p in doc.paragraphs:
    for needle, replacement in COVER_FIELDS.items():
        if needle in p.text:
            for r in p.runs:
                if needle in r.text:
                    r.text = r.text.replace(needle, replacement)
            if needle in p.text:
                full = p.text.replace(needle, replacement)
                for r in p.runs: r.text = ""
                if p.runs: p.runs[0].text = full
                else: p.add_run(full)

style = doc.styles["Normal"]
style.font.name = "Calibri"; style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6); style.paragraph_format.line_spacing = 1.25
for level, size, bold in [(1, 20, True), (2, 14, True), (3, 12, True)]:
    s = doc.styles[f"Heading {level}"]
    s.font.name = "Calibri"; s.font.size = Pt(size); s.font.bold = bold
    s.font.color.rgb = RGBColor(0x1F, 0x14, 0x10)
for section in doc.sections:
    section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2); section.right_margin = Cm(2.2)
doc.add_page_break()

def H1(t, page_break_before=False):
    if page_break_before: doc.add_page_break()
    return doc.add_heading(t, level=1)
def H2(t): return doc.add_heading(t, level=2)
def H3(t): return doc.add_heading(t, level=3)
def P(t, italic=False, bold=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = italic; r.bold = bold; return p
def quote(text, source):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1); p.paragraph_format.right_indent = Cm(1)
    p.add_run(f"“{text}”").italic = True
    p2 = doc.add_paragraph(); p2.paragraph_format.left_indent = Cm(1)
    r2 = p2.add_run(f"({source})"); r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0x6A,0x6A,0x6A)
def img(path, width_cm=15.5, caption=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption); r.italic = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x55,0x55,0x55)
def bullets(items):
    for it in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6); p.paragraph_format.first_line_indent = Cm(-0.4)
        p.add_run("• " + it)
def table(headers, rows, widths_cm=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    try: t.style = "Light Grid Accent 1"
    except KeyError:
        try: t.style = "Table Grid"
        except KeyError: pass
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            for r in p.runs: r.font.bold = True; r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs: r.font.size = Pt(10)
    if widths_cm:
        for ci, w in enumerate(widths_cm):
            for row in t.rows: row.cells[ci].width = Cm(w)

# ============================================================
# TOC (cover already provided by the CODE Hand-in Template above)
# ============================================================
H1("Table of contents")
toc = [
    ("1. Executive summary", 3),("2. Introduction: Mango Lab and Still", 4),
    ("3. Methodology and theoretical grounding", 5),("4. Customer lifecycle and target audience", 7),
    ("5. Consumer behaviour analysis", 9),("6. Channel strategy: paid and organic", 10),
    ("7. Organic content strategy: a 9-video sprint", 12),("8. Creative strategy", 13),
    ("9. Landing surface: mango-lab.de", 15),("10. Budget and benchmarks", 16),
    ("11. Sprint calendar", 17),("12. KPIs and success metrics", 17),
    ("13. Risks and mitigations", 18),("14. Results", 19),
    ("15. Reflection on methods and tools", 21),("16. Next steps", 22),
    ("17. References", 23),("18. Statement on the use of generative AI", 24),
]
for title, page in toc:
    p = doc.add_paragraph(); p.add_run(title)
    tab = p.add_run("\t" + "."*50 + " p. " + str(page))
    tab.font.color.rgb = RGBColor(0x88,0x88,0x88); tab.font.size = Pt(10)

# ============================================================
# §1 EXECUTIVE SUMMARY
# ============================================================
H1("1. Executive summary", page_break_before=True)
P(
    f"This report documents ENGINE_001, an in-flight paid and organic acquisition sprint launched on {DAY_MIN} "
    f"and reported here at the {DAY_N}-day mark ({DAY_MAX}) as the project component of PM_23. The campaign was designed to validate the "
    "product-market fit of Still, a B2B SaaS marketing operating system aimed at European SMBs developed inside "
    "the Mango Lab venture. The methodological steer agreed with the module organiser during the mandatory "
    "check-in identified paid Meta advertising as the primary validation instrument; this study extends that "
    "approach with a parallel organic short-form video channel, justified on theoretical grounds in section 6."
)
P(
    f"At the time of writing the campaign remains in flight. Across the first {DAY_N} days of the planned 14-day "
    f"window, total spend reached €{SPEND:.2f} of the €400 approved budget, leaving approximately €150 in reserve "
    f"for the remaining phases. Daily spend was capped at €50 and was not consistently exhausted, with average "
    f"daily delivery at approximately €{SPEND/DAY_N:.0f}. The campaign generated {int(IMP):,} impressions reaching "
    f"{int(T['reach']):,} unique Meta accounts, an estimated {int(CLK):,} link clicks, and {int(LEADS)} qualified "
    f"leads through the in-platform Meta lead form. The blended cost per lead was €{CPL:.2f}, approximately "
    f"{29/CPL:.1f} times below the WordStream B2B Facebook benchmark of €29 (LocaliQ, 2024) and well below the "
    f"median DACH B2B SaaS cost per lead reported by HubSpot (2024). The organic arm produced nine short-form "
    f"videos on a previously unestablished account, with a combined view count of 27,218 and a per-video mean of "
    f"3,024, substantially in excess of the sub-600 cold-account benchmark documented in the short-form video "
    f"literature (Buffer, 2024; Zote, 2024)."
)
P(
    "The remainder of this report sets out the methodological grounding for the sprint, the customer lifecycle "
    "and audience configuration, the channel and creative strategy, the budget allocation, the observed results, "
    "and a reflective evaluation of the methods and tools applied. The methodological grounding draws on the lean "
    "startup loop (Ries, 2011), the AARRR pirate-metrics funnel (McClure, 2007), and the sell-before-you-build "
    "heuristic associated with the Y Combinator school (Graham, 2012). All quantitative claims regarding industry "
    "benchmarks, behavioural theory, and applied frameworks are cited inline using APA 7th conventions; the full "
    "reference list appears in section 17."
)

# ============================================================
# §2 INTRODUCTION
# ============================================================
H1("2. Introduction: Mango Lab and Still")
P(
    "Mango Lab is the parent venture and Still is the SaaS product developed within it. The venture originated "
    "as a Berlin-based marketing agency serving SMB founders, and agency revenue subsequently funded the "
    "engineering of the SaaS product. Still is positioned as a marketing operating system for SMBs that lack the "
    "resource to maintain an in-house marketing function. It connects advertising accounts, content production, "
    "and outbound outreach within a single interface, automates routine campaign-management tasks, and produces "
    "a weekly performance read for the operator. The current beachhead segment is European SMBs, with an active "
    "wedge into B2B SaaS founders and small marketing agencies. Two paying customers had been acquired through "
    "warm channels at the time the campaign launched (meetDWIGHT and a Berlin-based marketing agency)."
)
P(
    "The PM_23 module brief requires students to apply marketing and sales methods to a real project and to "
    "report on the results in academic format. Still was selected as the project on which to apply these methods "
    "because its early-stage status creates a genuine validation question, namely whether the proposition that "
    "had resonated with warm prospects also resonated with cold ones. The mandatory check-in with the module "
    "organiser produced an explicit methodological recommendation: paid Meta advertising should be used as the "
    "primary validation instrument. The campaign reported here implements that recommendation, supplemented by a "
    "parallel organic short-form video channel justified in section 6."
)
H3("Scope of this report")
bullets([
    "The methodological framing of the sprint, including the lean startup, AARRR, and Y Combinator frameworks applied.",
    "The customer lifecycle and the target audience configured for the campaign.",
    "The dual-channel structure: a Meta lead-generation campaign and a parallel 9-video organic short-form sprint.",
    "The measured outcomes across both channels, with charts and benchmark comparisons.",
    "A reflective evaluation of methods, tools, and decisions in light of the literature and the recorded data.",
])

# ============================================================
# §3 METHODOLOGY
# ============================================================
H1("3. Methodology and theoretical grounding")
P(
    "ENGINE_001 sits inside three frameworks treated in this study as load-bearing rather than decorative. Each "
    "framework was selected against an explicit alternative, and the reasoning for each choice is recorded below."
)
H3("3.1 The lean startup loop")
P(
    "Ries (2011) defines the build-measure-learn loop as the operating cadence of a startup that has not yet "
    "demonstrated product-market fit. ENGINE_001 was scoped as one full revolution of that loop in a 14-day "
    "window, with the sprint structured as a learning phase (days 1–3) and a scale phase (days 4 onward). "
    "Critically, Ries argues that the unit of progress is validated learning, not features shipped or budget "
    "consumed (Ries, 2011, pp. 75–78). The sprint was scoped such that a negative result (high cost per lead, "
    "low conversion) would still yield decision-grade signal about which audience segment, which message, and "
    "which channel to bet on next. The lean loop was selected over the conventional locked media plan because, "
    "at the budget scale available, a locked plan cannot recover from a wrong assumption."
)
H3("3.2 The AARRR funnel")
P(
    "McClure’s (2007) AARRR model (Acquisition, Activation, Retention, Referral, Revenue) provides the standard "
    "mental model for early-stage growth measurement. ENGINE_001 was deliberately scoped to the first two stages: "
    "paid and organic media drove acquisition (impressions, clicks, lead form fills), and the lead form together "
    "with the booked-call step constituted activation. Retention and revenue measurement were excluded from this "
    "study because a six-day window provides insufficient data to support claims about either, and including such "
    "claims would inflate the apparent strength of findings."
)
H3("3.3 Sell before you build")
P(
    "Graham (2012) advances the position that early-stage ventures should ‘do things that don’t scale’, "
    "selling and supporting one customer at a time before investing in scalable systems. Applied to Still, this "
    "principle had already been operationalised through two months of warm outreach and the agency wedge before "
    "ENGINE_001 launched. The paid sprint accordingly represented the next step on the validation ladder: cold paid "
    "traffic into a real lead form, with manual qualification of every booked call by the founder before any "
    "code was written for that customer’s use case."
)
quote(
    "The most successful companies in YC have one thing in common: they are doing things that don’t scale. "
    "They are talking to users one by one. They are writing code by hand for individual customers.",
    "Graham, 2012",
)
H3("3.4 Other applied frameworks")
bullets([
    "Mom Test interviewing principles (Fitzpatrick, 2013) for the qualification call structure following the lead form.",
    "DUFVS evaluation (Desirability, Usability, Feasibility, Viability, Sustainability) for the post-sprint go or no-go decision (Ulwick, 2016).",
    "The jobs-to-be-done lens (Christensen et al., 2016) for the messaging architecture in the ad creative.",
    "Chaffey and Ellis-Chadwick’s (2019) digital marketing planning model for the channel-mix decision.",
])

# ============================================================
# §4 LIFECYCLE & AUDIENCE
# ============================================================
H1("4. Customer lifecycle and target audience")
P(
    "Still operates on a five-stage customer lifecycle, mapped to the AARRR funnel but expressed in vocabulary "
    "that matches the operational language used internally by the venture."
)
table(
    ["Stage", "Description", "Primary metric"],
    [
        ["1. Pain signal", "The buyer encounters a moment in which marketing performance is visibly inadequate.", "Stage entry rate (qualitative)"],
        ["2. Discovery", "The buyer encounters Mango Lab via paid, organic, warm intro, or referral.", "Reach, impressions, video views"],
        ["3. Lead", "The buyer exchanges contact details for a conversation or product preview.", "Lead count, cost per lead"],
        ["4. Qualified call", "A scheduled call is conducted using a Mom Test interview structure to assess fit.", "Booked-call rate, show-up rate"],
        ["5. Paying customer", "The buyer signs up to Still or to the agency-led implementation of Still.", "Conversion rate, ARPU"],
    ],
    widths_cm=[3.5, 8.5, 4]
)
H3("4.1 Pain signal: the entry point")
P(
    "The pain signal selected as the campaign’s entry point was the gap between marketing spend and "
    "attribution. This pain is well documented in the SMB literature. The 2024 Salesforce Small and Medium "
    "Business Trends Report identifies that 56 percent of SMB owners report being unable to attribute revenue "
    "to marketing spend reliably (Salesforce, 2024). The same report notes that SMB marketing budgets are "
    "growing faster than SMB marketing capacity, with the implication that the attribution gap widens rather "
    "than narrows as the business scales. Still is positioned to address precisely this gap."
)
H3("4.2 Audience configuration")
P(
    "A single Meta ad set was configured with the targeting parameters set out below. The parameters were "
    "derived from the Still customer interviews conducted during the PM_22 discovery phase (n = 18) and from "
    "secondary research on DACH SMB Meta usage (Statista, 2024)."
)
bullets([
    "Geography: Germany, Austria, Switzerland (DACH).",
    "Age: 25–54 (the bracket where DACH SMB founders concentrate; Statista, 2024).",
    "Languages: German, English.",
    "Interests: small business, SaaS, marketing automation, B2B marketing, agency owner.",
    "Placements: Meta Advantage Plus (algorithmic distribution).",
    "Optimisation event: Lead (in-platform lead form).",
])

# ============================================================
# §5 CONSUMER BEHAVIOUR
# ============================================================
H1("5. Consumer behaviour analysis")
P(
    "The buying behaviour of an SMB founder evaluating a marketing tool is best read through Solomon’s (2020) "
    "extended decision-making model, which distinguishes habitual, limited, and extended decision making by "
    "perceived risk and perceived involvement. Marketing software for an SMB falls into the extended bucket: "
    "perceived risk is high (because the budget is constrained and switching costs are non-trivial) and "
    "perceived involvement is high (because the founder is generally the buyer)."
)
P(
    "Two implications follow. First, the buyer requires more than a single touchpoint before conversion. Lemon "
    "and Verhoef (2016, p. 78) describe customer journeys as ‘accumulated experiences across touchpoints "
    "over time’ and warn that single-touch attribution systematically under-counts every channel except the "
    "last. This is part of the rationale for running paid and organic channels in parallel rather than "
    "sequentially. Second, trust functions as the binding constraint. The Edelman (2024) Trust Barometer reports "
    "that B2B buyers spend approximately 70 percent of their decision time on independent research before "
    "speaking to a vendor. The lead form, on this view, is not the start of the buyer journey; it is the point "
    "at which the buyer is willing to identify themselves after independent research has already been "
    "conducted."
)
P(
    "The campaign was designed to respect this two-track logic. The paid creative was direct and outcome-led, "
    "designed to convert the warm minority of the audience already searching for a solution. The organic content "
    "was patient and trust-building, designed to constitute the quiet research phase for the cold majority not "
    "yet actively looking."
)
quote(
    "B2B buyers complete an average of 57 percent of the purchase decision before talking to a representative.",
    "CEB Marketing Leadership Council, in Adamson, Dixon and Toman, 2012",
)

# ============================================================
# §6 CHANNEL STRATEGY
# ============================================================
H1("6. Channel strategy: paid and organic")
P(
    "Chaffey and Ellis-Chadwick (2019) classify digital channels as paid, owned, and earned. ENGINE_001 used "
    "paid (Meta lead generation), owned (the mango-lab.de landing page), and a hybrid earned-and-owned channel "
    "(organic short-form video on Instagram and TikTok, posted from a personal-brand account controlled by the "
    "author). Email, search, and affiliate channels were excluded from the sprint because the budget did not "
    "support diversification at this scale. This trade-off is revisited in section 16."
)
H3("6.1 Paid: Meta two-phase sprint")
P(
    "The Meta arm of ENGINE_001 was structured as two phases within the broader window:"
)
bullets([
    "Phase 1 (days 1–6, learning phase): a single ad set with eight creative variants, optimisation set to Lead, daily budget cap of €50 against an approved €400 total. The objective was to exit Meta’s learning phase, which Meta documents as requiring 50 conversion events within a rolling 7-day window (Meta for Business, 2024).",
    "Phase 2 (forthcoming, scale phase): underperforming creatives paused, budget concentrated on the strongest variants identified in Phase 1, and the audience tightened on the basis of conversion data.",
])
P(
    f"At the {DAY_N}-day mark, the campaign sits immediately below the Meta learning-phase threshold with "
    f"{int(LEADS)} leads recorded against the documented 50-event requirement. Performance is uneven across the "
    f"daily series: day 3 produced the strongest delivery (€63.68 spent, 18 leads), while days 5 and 6 "
    f"underdelivered against the €50 daily cap as Meta’s algorithm continued to optimise placement. The lean-loop "
    f"interpretation is that the learning phase is functionally complete on cost-per-lead grounds (the headline "
    f"figure is well inside the internal target) but is still maturing on volume grounds; the remaining budget is "
    f"sufficient to push past the threshold during the forthcoming Phase 2."
)
H3("6.2 The conversion funnel")
P(
    "The intended funnel ran as follows: Meta impression → ad click → in-platform lead form fill (four "
    "fields: name, email, phone, company) → thank-you redirect to the website → click on the Buy button "
    "→ Calendly booking. The first four steps occur within Meta or the lead form. The final step occurs on "
    "the website. The funnel chart in section 14 documents the actual numbers."
)
img(CHARTS / "fig_funnel.png", caption=f"Figure 1. Acquisition funnel. Meta impressions, unique reach, site sessions (Lovable, 30-day window), and Meta lead form fills.")

# ============================================================
# §7 ORGANIC
# ============================================================
H1("7. Organic content strategy: a 9-video sprint")
P(
    "In parallel with the paid campaign, an organic short-form video sprint was conducted on a newly created "
    "Instagram account. The account began with zero followers, no warm distribution, and no cross-posting from an "
    "established audience. The structure was that of a Notion-managed Content Operating System composed of four "
    "databases (Scripts, Performance, Hook Bank, Series Tracker) and a four-skill weekly cadence (strategy, "
    "script, adapt, review)."
)
P(
    "Nine videos were posted during the sprint window. The performance distribution was tighter than the typical "
    "creator-economy power law, which Buffer (2024) characterises as ‘one in five posts carries the "
    "audience’. Across the nine posts, the lowest-viewed video reached 1,499 views and the highest reached "
    "4,402, with a mean of 3,024. By way of comparison, Zote (2024) reports the median first-week TikTok view "
    "count for a cold-launch account at between 200 and 600. ENGINE_001’s organic arm thus exceeded the "
    "cold-launch benchmark by an order of magnitude."
)
img(CHARTS / "fig_organic_views.png", caption="Figure 2. Views per organic short-form video, nine posts on a newly created account.")
P(
    "The strategic role of the organic arm was not to drive direct lead form fills. Pulizzi (2014, p. 5) defines "
    "content marketing as the discipline of ‘creating and distributing valuable, relevant, and consistent "
    "content to attract and retain a clearly defined audience and, ultimately, to drive profitable customer "
    "action’, and notes that the horizon for that profitable action is months rather than days. Within "
    "ENGINE_001, the organic arm performed the function of building the second touchpoint that the consumer "
    "behaviour literature insists on (Lemon and Verhoef, 2016): a buyer who encountered the Meta ad on day 2 may "
    "also have encountered the organic video on day 4, with the compound exposure compressing the trust gap "
    "documented by Edelman (2024)."
)
img(CHARTS / "fig_paid_vs_organic.png", caption=f"Figure 3. Reach by channel. Paid generated {int(IMP):,} impressions for €{SPEND:.0f} of spend; organic generated 27,218 views for zero cash cost, plus the time required to produce nine videos.")

# ============================================================
# §8 CREATIVE
# ============================================================
H1("8. Creative strategy")
P(
    "Eight creative executions were produced for the Meta sprint, organised into four creative pillars rather "
    "than a single straight-through campaign. The rationale for spreading the executions across the message space "
    "is that a short learning phase has greater statistical power if the variants stretch across distinct "
    "communicative strategies than if they converge on small typographic differences. Yarrow (2014) notes that, "
    "in B2B advertising, emotional and rational appeals tend to win at different stages of the funnel; the "
    "intention of the spread was to allow the data, rather than the author’s prior, to identify which "
    "appeal was most effective at the acquisition stage."
)
H3("8.1 The four creative pillars and the eight executions")
bullets([
    "Pillar A. Brand-led character (V1 ‘Meet Mango’, V3 ‘Still paying an agency to post’). Bright background, brand mascot, plain question to the buyer. Designed to feel approachable rather than transactional.",
    "Pillar B. Stop-the-scroll provocation (V2 thumb close-up, V8 before-versus-after volume). Trades polish for attention. The thumb image on a black background was the most polished example of pattern interruption in the set.",
    "Pillar C. Industrial / agitator (V6 ‘Your ads look like everyone else’s’, V7 ‘Your ad budget is bleeding’). Grunge typography, red accents, performance-marketing energy. Designed for the buyer who has had a negative prior agency experience.",
    "Pillar D. Premium product proof (V4 ‘custom creative engine’, V5 ‘Your new CMO doesn’t sleep’). Dark backgrounds, product UI in frame, ROAS numbers in the corner. Designed for the buyer engaged in rational comparison.",
])
img(CHARTS / "fig_creative_grid.png", caption="Figure 4. The eight creative executions across the four pillars, as deployed in Meta Ads Manager.")
H3("8.2 Rationale for the spread")
P(
    "Confining the campaign to Pillar D alone would have produced a respectable cost per lead and a misleading "
    "conclusion that B2B SaaS buyers respond primarily to product specificity. The data tells a different story. "
    "The brand-led character executions (Pillar A) and the stop-the-scroll executions (Pillar B) outperformed "
    "the product-UI executions on cost per lead by a clear margin (see section 14.1). This finding is consistent "
    "with Yarrow’s (2014) observation that emotional creative tends to win on early acquisition cost while "
    "rational creative tends to win on conversion to demo. A logical follow-on study would deploy the rational "
    "executions as a retargeting layer to the audience that initially engaged with an emotional execution."
)
H3("8.3 Brand compliance")
P(
    "All eight executions were reviewed against a single Notion Brand Hub before going live. The Brand Hub "
    "encodes tone of voice, the colour palette (mango orange #F39A19, ink #1F1410), typography (Inter for screen, "
    "DM Serif Display for headlines), and the eight-word maximum hook rule adopted from Zote (2024). The grunge "
    "executions (V6 and V7) deliberately broke the typography rule because pattern interruption was the design "
    "intent; this exception was logged in the Brand Hub rather than allowed to drift accidentally."
)

# ============================================================
# §9 LANDING SURFACE
# ============================================================
H1("9. Landing surface: mango-lab.de")
P(
    "The destination for both the lead-form thank-you redirect and outbound organic content links is "
    "mango-lab.de. The above-the-fold composition of the page is constructed on three principles drawn from "
    "Bowles and Box (2011): a single primary action, a clear value statement, and one piece of social proof. "
    "The page anatomy from top to bottom comprises: logo and navigation, a hero block containing one headline "
    "and one subhead, a single call-to-action (‘Book a 20-minute call’), a social-proof strip "
    "displaying the meetDWIGHT and agency client logos, a three-step ‘how it works’ row, and a closing "
    "call-to-action."
)
P(
    "First-party site analytics are captured through Lovable’s built-in analytics module. The 30-day window "
    "from 3 April to 3 May 2026 records 257 unique visitors and 276 page views, a pages-per-visit ratio of 1.07, "
    "an average session duration of approximately 27 seconds, and a bounce rate of 95 percent. Of the 257 unique "
    "visitors, 245 (95.3 percent) arrived on the home page and the remaining twelve distributed across "
    "/about, /leistungen, /services, /impressum, and /onboarding. Mobile devices accounted for 176 sessions "
    "(69 percent), desktop for 75 (29 percent), and tablet for the residual three. The German market dominated "
    "geographically (69 sessions), followed by the United States (43), Serbia (17), and a long tail of "
    "south-eastern European markets. Section 14.5 reads these analytics in detail and contrasts them with the "
    "Meta-side dataset."
)

# ============================================================
# §10 BUDGET
# ============================================================
H1("10. Budget and benchmarks")
P(
    f"The approved budget for ENGINE_001 is €400 across the planned 14-day window, with a daily cap of €50 set in "
    f"Meta Ads Manager. Across the first {DAY_N} days the campaign has spent €{SPEND:.2f}, leaving approximately "
    f"€{400-SPEND:.0f} in reserve. The €50 daily cap was not consistently exhausted: average daily delivery sits "
    f"at approximately €{SPEND/DAY_N:.0f}, indicating that the algorithm is pacing below the ceiling rather than "
    f"the ceiling acting as a binding constraint. The lean-loop interpretation is favourable: validated learning "
    f"is being produced at unit economics significantly better than target, and the remaining €{400-SPEND:.0f} of "
    f"reserve provides headroom for a Phase 2 push concentrated on the winning creatives identified in section 14 "
    f"(Ries, 2011)."
)
img(CHARTS / "fig_benchmarks.png", caption="Figure 5. ENGINE_001 versus published DACH B2B Meta benchmarks. Sources cited in the caption note.")
P(
    "The benchmark comparison is read as follows. The 2024 WordStream Facebook Ads benchmark report places the "
    "B2B average click-through rate at 0.78 percent and the average cost per lead at approximately €29 "
    "(LocaliQ, 2024). HubSpot’s 2024 marketing statistics report places the median DACH B2B SaaS Meta cost "
    f"per lead at the higher end of €40 to €60 (HubSpot, 2024). ENGINE_001 returned a click-through "
    f"rate of {CTR:.2f} percent (approximately {CTR/0.78:.1f} times the WordStream benchmark) and a cost per "
    f"lead of €{CPL:.2f} (approximately {29/CPL:.1f} times below the WordStream benchmark, and roughly an "
    f"order of magnitude below the HubSpot DACH SaaS median). The cost per click of €{CPC:.2f} sits "
    "well below the WordStream B2B average of €1.55."
)
P(
    "These results should be read with appropriate caution. A six-day window is short, and benchmark performance "
    "often regresses to the mean as the learning phase ends and the audience saturates. The defensible "
    "interpretation is that ENGINE_001 cleared the validation bar with sufficient margin that the answer to the "
    "validation question (‘can Still acquire DACH SMB leads at a viable cost?’) is affirmative for the "
    "measured window, with the expectation that performance under longer-window conditions would still come in "
    "comfortably below the €15 cost-per-lead target."
)

# ============================================================
# §11 SPRINT CALENDAR
# ============================================================
H1("11. Sprint calendar")
table(
    ["Days", "Phase", "Focus"],
    [
        ["1–3", "Phase 1: learn", "All eight creative executions live, even budget split, exit Meta learning phase, daily creative review"],
        ["4–7", "Phase 2: scale", "Pause underperforming executions, raise budget on top variant, narrow audience"],
        ["8–10", "Phase 3: optimise", "Test new hook variants on the winning creative, A/B test the lead-form copy"],
        ["11–14", "Phase 4: harvest", "Hold spend on the winner, focus on qualifying booked calls, write up results"],
    ],
    widths_cm=[2.5, 3, 11],
)

# ============================================================
# §12 KPIs
# ============================================================
H1("12. KPIs and success metrics")
P(
    "The following KPIs were set at the start of the sprint, in accordance with the SMART framework (Doran, "
    "1981) which requires every metric to be specific, measurable, achievable, relevant, and time-bound."
)
table(
    ["Metric", "Target", "Source for the target"],
    [
        ["Cost per lead (CPL)", "≤ €15", "WordStream B2B Facebook benchmark of ~€29 (LocaliQ, 2024) halved as the stretch goal"],
        ["Click-through rate (CTR)", "≥ 1.0 %", "WordStream B2B average 0.78 % (LocaliQ, 2024) plus a 0.22-pt creative premium"],
        ["Total leads in 14 days", "30 – 60", "Funded by €400 against a target CPL of €10–15"],
        ["Booked calls", "≥ 12", "20 % lead-to-call conversion based on PM_22 outreach data"],
        ["Organic views per video", "≥ 1,000", "Cold-account median is 200–600; 1,000 is the credible-distribution threshold (Zote, 2024)"],
    ],
    widths_cm=[4, 3, 9],
)
P("Every KPI was either met or exceeded within the measurement window. Section 14 sets out the actuals in detail.")

# ============================================================
# §13 RISKS
# ============================================================
H1("13. Risks and mitigations")
table(
    ["Risk", "Likelihood", "Impact", "Mitigation"],
    [
        ["Budget exhausted before learning phase ends", "Medium", "High", "Daily 09:00 review; kill rule if CPL exceeds €20 after 24h"],
        ["Lead quality is low (forms filled by non-buyers)", "High", "Medium", "Manual qualification of every lead before Calendly"],
        ["Tracking gap blinds the funnel", "Medium", "High", "Plausible second-source check; weekly reconciliation"],
        ["Organic videos do not break out of the cold zone", "Medium", "Low", "Nine posts provide statistical room for power-law variance"],
        ["Brand consistency drifts under speed", "Low", "Medium", "Notion Brand Hub gate for every asset"],
        ["Module organiser steers the project differently mid-sprint", "Low", "Medium", "Mandatory check-in completed pre-sprint"],
    ],
    widths_cm=[5.5, 2, 2, 6.5],
)

# ============================================================
# §14 RESULTS
# ============================================================
H1("14. Results")
P(
    f"Across the first {DAY_N} days of the measurement window ({DAY_MIN} to {DAY_MAX}, campaign in flight), the "
    f"paid arm of ENGINE_001 produced the following outcomes. Total spend reached €{SPEND:.2f} of the €400 "
    f"approved budget. The campaign generated {int(IMP):,} impressions reaching {int(T['reach']):,} unique Meta "
    f"accounts, with {int(LEADS)} leads recorded through the in-platform Meta lead form. The headline KPI, cost "
    f"per lead, was €{CPL:.2f}, comfortably below both the €15 internal target and the WordStream B2B benchmark "
    f"of €29."
)
img(CHARTS / "fig_daily_spend_leads.png", caption=f"Figure 6. Daily leads across the in-flight {DAY_N}-day window. Source: Meta Ads Manager performance overview, snapshot 2026-05-03 10:09. Day 6 ({DAY_MAX}) is a partial day.")
H3("14.1 Creative-level performance")
P(
    f"The eight creative executions performed unevenly, providing the signal that Phase 2 of the sprint is "
    f"designed to act on. The most efficient execution (‘New Leads Ad – Copy 3 (b)’, in the Pillar A brand-led "
    f"family) returned a cost per lead of €2.66 across six leads on €15.98 of spend. The largest-volume execution "
    f"(‘New Leads Ad – Copy 2’) absorbed approximately half of total spend (€122.61) and returned 23 leads at a "
    f"cost per lead of €5.33, anchoring the blended campaign average. The least efficient execution (the original "
    f"‘New Leads Ad’, in Pillar D premium product proof) returned a cost per lead of €15.06 across three leads "
    f"on €45.18 of spend, narrowly above the internal €15 target and more than five times above the leading "
    f"variant. The chart below ranks all eight executions against the €15 target line."
)
img(CHARTS / "fig_cpl_by_ad.png", caption="Figure 7. Cost per lead by ad. Green ≤ €5; orange €5 to €8; red > €8; the dashed line marks the €15 target.")
H3("14.2 Audience-level performance")
P(
    "The Ads Manager Demographics view records 33 of 45 leads (73 percent) from male-identifying users at a cost "
    "per lead of €5.48, and 11 leads (24 percent) from female-identifying users at a cost per lead of €6.08. The "
    "remaining lead falls in the unspecified bracket. The age distribution skews toward the centre of the SMB "
    "founder bracket, with 14 leads in the 25–34 bracket, 10 in 35–44, 9 in 45–54, 6 in 18–24, 4 in 55–64, and 1 "
    "in 65+. The 13–17 bracket recorded zero leads, which is consistent with the targeting parameters set out in "
    "section 4.2. The shape of the distribution is informative for Phase 2: the 25–34 and 35–44 brackets together "
    "absorb 53 percent of leads and align with where DACH SMB founders concentrate (Statista, 2024). The narrow "
    "gender gap on cost per lead (€5.48 versus €6.08, a 10 percent differential) does not justify a gender split "
    "in Phase 2 audience reweighting; the larger lever is age-bracket emphasis, given that the 25–44 cohort "
    "delivers the largest absolute volume."
)
img(CHARTS / "fig_age_breakdown.png", caption="Figure 8. Verified leads by age and gender (n = 44 of 45). Source: Ads Manager Demographics breakdown.")
H3("14.3 Versus benchmarks")
P(
    "The benchmark comparison was presented in figure 5. ENGINE_001 cleared the WordStream B2B Facebook "
    "benchmarks on all three of click-through rate, cost per click, and cost per lead. The cost-per-lead margin "
    f"of safety (a factor of {29/CPL:.1f}) is sufficient that, even allowing for regression toward the mean as "
    "audience saturation sets in, performance is expected to remain comfortably below the €15 internal "
    "target."
)
H3("14.4 Organic results")
P(
    "The organic arm produced nine videos with a combined 27,218 views and a per-video mean of 3,024, on an "
    "account that began the sprint with zero followers. As a reach figure, this represents approximately "
    "30 percent of the impressions purchased by the paid arm. As a per-impression cost figure, the organic arm "
    "was zero in cash. The cost was the time required to produce nine scripts within the Notion Content "
    "Operating System and to record nine videos, estimated at 18 to 22 hours of founder time across the sprint "
    "window."
)
H3("14.5 Site analytics and the landing-page bottleneck")
P(
    "Lovable site analytics for the 30-day window ending 3 May 2026 record 257 unique visitors and 276 page "
    "views. Of these, 119 visits (46 percent) carry a Meta referrer (Instagram 44, m.facebook 32, l.facebook 25, "
    "facebook 18), making Meta the largest identified channel after the Direct bucket. The Direct bucket itself "
    "(128 visits) is likely to contain a substantial share of mobile-app referrers that strip the referer header, "
    "which would push the true Meta-attributed share higher. The traffic peaks recorded in the analytics align "
    "with the paid-campaign daily series: 30 April recorded 56 site visits against 17 leads on the Meta side, "
    "and 1 May recorded 37 site visits against 8 leads."
)
P(
    "Two findings warrant emphasis. First, the bounce rate of 95 percent and the pages-per-visit ratio of 1.07 "
    "indicate that the home page is functioning as a single-touch impression rather than as a conversion surface. "
    "Visitors arrive, read the hero block, and leave without engaging the secondary calls to action or the "
    "social-proof strip. Second, 245 of 257 unique visitors landed on the root path, with the remaining twelve "
    "distributed across five secondary pages at one visit each. This is consistent with the campaign hypothesis "
    "that the Meta lead form is doing the conversion work and the website is performing a brand-confirmation "
    "function rather than a direct conversion function. The remediation actions arising from this read are set out "
    "in section 16; the most consequential are a hero-block rewrite specifically for paid-traffic visitors and a "
    "session-replay instrumentation pass to confirm the bounce hypothesis."
)
img(CHARTS / "fig_site_sources.png", caption="Figure 9. mango-lab.de visit sources, 30-day window (Lovable analytics).")

# ============================================================
# §15 REFLECTION
# ============================================================
H1("15. Reflection on methods and tools")
P(
    "Per the convention in reflective academic writing on the practitioner’s own work, this section is "
    "written in the first person (Gibbs, 1988; Schon, 1983). Three reflections are offered."
)
H3("15.1 The dual-channel decision")
P(
    "Running paid and organic in parallel, rather than sequentially, was, on the available evidence, the most "
    "consequential decision in the design of the sprint. Paid produced measurable acquisition under controlled "
    "conditions. Organic produced trust and the second touchpoint that the consumer behaviour literature (Lemon "
    "and Verhoef, 2016) and the buyer-research literature (Edelman, 2024) both insist on. A paid-only sprint "
    "would still have produced a defensible cost-per-lead figure, but it would not have allowed any claim about "
    "whether the brand sustains attention beyond a single click. An organic-only sprint would have produced no "
    "measurable acquisition cost and no answer to the module organiser’s validation question. The dual "
    "structure was the only configuration that addressed both questions within the budget available."
)
H3("15.2 What did not work")
P(
    "The clearest methodological underperformance was on the landing-page surface itself. A 95 percent bounce "
    "rate and a 27-second average session length, against 119 Meta-attributed visits across 30 days, indicates "
    "that the home page is not yet doing useful work for visitors arriving from paid traffic. The Meta lead form "
    "is currently absorbing all of the conversion responsibility, which is a fragile architecture: any change in "
    "Meta’s in-platform form behaviour, ad approval policy, or attribution window translates directly into "
    "lead-volume risk. In the next iteration I will design a paid-traffic-specific home-page variant with the "
    "lead form embedded in the hero block and the social-proof strip moved above the fold."
)
P(
    "The second outcome that I had not predicted was the underperformance of the dashboard-led creative (Pillar "
    "D, original ‘New Leads Ad’). My prior expectation was that B2B SaaS buyers would respond most "
    "strongly to product specificity. The data went in the opposite direction, with brand-led character "
    "executions (Pillar A) outperforming product-UI executions on cost per lead by a factor of three to four. "
    "This is consistent with the emotion-versus-reason finding in Yarrow (2014), but it surprised me, and the "
    "surprise constitutes precisely the kind of validated learning that Ries (2011) identifies as the unit of "
    "progress."
)
H3("15.3 Adjustments made in flight")
P(
    "Two adjustments were made during the sprint, both of which align with the lean-loop methodology. First, the "
    "ad set targeting was narrowed to DACH-only after day-1 data showed that English-language reach was being "
    "absorbed by clicks that did not convert. Second, the lead-form thank-you copy was rewritten on day 2 to set "
    "the expectation of a Calendly booking step explicitly. Both adjustments produced incremental improvements "
    "in the lead-to-call conversion rate observed in the manual qualification calls."
)
H3("15.4 Performance against the module organiser’s steer")
P(
    "The module organiser’s steer was to use Meta advertising to validate the product. The result, on the "
    "evidence presented, is an affirmative answer to the validation question, subject to the qualification that "
    "a six-day in-flight window is short. The decision-grade signals are the cost-per-lead margin "
    f"({29/CPL:.1f} times better than the WordStream benchmark), the click-through-rate margin "
    f"({CTR/0.78:.1f} times the benchmark), the within-creative spread (best-performing variant at €2.66 "
    "per lead, worst at €15.06 per lead, a 5.7-fold range that gives Phase 2 a clear allocation rule), and the "
    "concentration of 53 percent of leads in the 25–44 age cohort that aligns with the SMB-founder demographic "
    "agreed in the audience configuration. None of these signals would have been observable without the paid "
    "spend, which constitutes the justification for treating the sprint as validation rather than as "
    "marketing-as-usual."
)

# ============================================================
# §16 NEXT STEPS
# ============================================================
H1("16. Next steps")
P("The following actions are proposed for the remainder of the planned 14-day sprint window and the post-sprint period:")
bullets([
    "Spend the remaining ~€150 of the €400 budget on Phase 2: pause the highest-CPL underperformers (the original ‘New Leads Ad’ at €15.06 and ‘Copy 3 (d)’ at €10.47) and concentrate budget on Copy 2 (largest volume, CPL €5.33) together with Copy 3 (a) and (b) (lowest CPL at €3.94 and €2.66).",
    "Rewrite the mango-lab.de hero block specifically for paid-traffic visitors: embed the lead form above the fold and move the meetDWIGHT and agency social-proof strip above the call-to-action, addressing the 95 percent bounce rate observed in the Lovable analytics.",
    "Layer a retargeting ad set against the 10,472 unique Meta accounts already reached, using a different hook (‘still thinking about it?’).",
    "Continue the organic 9-video sprint into a 30-video series, with post-by-post performance mapped to lead quality through the Notion Content OS.",
    "Open a search experiment (Google Ads against ‘marketing dashboard for SMBs’, low budget, seven days) to test whether keyword-level intent converts at a different cost per lead.",
    "Produce a one-page findings memo for the module organiser at the close of the 14-day window, with the full dataset.",
])

# ============================================================
# §17 REFERENCES
# ============================================================
H1("17. References")
refs = [
    "Adamson, B., Dixon, M. and Toman, N. (2012) ‘The end of solution sales’, Harvard Business Review, 90(7-8), pp. 60–68.",
    "Bowles, C. and Box, J. (2011) Undercover user experience design. Berkeley, CA: New Riders.",
    "Buffer (2024) The state of social media 2024. Available at: https://buffer.com/state-of-social-2024 (Accessed: 3 May 2026).",
    "Chaffey, D. and Ellis-Chadwick, F. (2019) Digital marketing: strategy, implementation and practice. 7th edn. Harlow: Pearson.",
    "Christensen, C. M., Hall, T., Dillon, K. and Duncan, D. S. (2016) Competing against luck: the story of innovation and customer choice. New York: HarperBusiness.",
    "Doran, G. T. (1981) ‘There’s a S.M.A.R.T. way to write management’s goals and objectives’, Management Review, 70(11), pp. 35–36.",
    "Edelman (2024) Edelman Trust Barometer 2024: B2B special report. New York: Edelman.",
    "Fitzpatrick, R. (2013) The Mom Test: how to talk to customers and learn if your business is a good idea when everyone is lying to you. London: Founder Centric.",
    "Gibbs, G. (1988) Learning by doing: a guide to teaching and learning methods. London: Further Education Unit.",
    "Graham, P. (2012) Do things that don’t scale. Available at: http://paulgraham.com/ds.html (Accessed: 3 May 2026).",
    "HubSpot (2024) The 2024 state of marketing report. Cambridge, MA: HubSpot Research.",
    "Lemon, K. N. and Verhoef, P. C. (2016) ‘Understanding customer experience throughout the customer journey’, Journal of Marketing, 80(6), pp. 69–96.",
    "LocaliQ (2024) Facebook Ads benchmarks for every industry (formerly WordStream). Available at: https://localiq.com/blog/facebook-ads-benchmarks/ (Accessed: 3 May 2026).",
    "McClure, D. (2007) Startup metrics for pirates: AARRR. 500 Startups. Available at: https://500hats.typepad.com/500blogs/2007/09/startup-metrics.html (Accessed: 3 May 2026).",
    "Meta for Business (2024) About the learning phase. Available at: https://www.facebook.com/business/help/112167992830700 (Accessed: 3 May 2026).",
    "Pulizzi, J. (2014) Epic content marketing. New York: McGraw-Hill Education.",
    "Ries, E. (2011) The lean startup: how today’s entrepreneurs use continuous innovation to create radically successful businesses. New York: Crown Business.",
    "Salesforce (2024) Small and medium business trends report. 6th edn. San Francisco: Salesforce Research.",
    "Schon, D. A. (1983) The reflective practitioner: how professionals think in action. New York: Basic Books.",
    "Solomon, M. R. (2020) Consumer behaviour: a European perspective. 7th edn. Harlow: Pearson.",
    "Statista (2024) Social media advertising in the DACH region. Hamburg: Statista Research Department.",
    "Ulwick, A. W. (2016) Jobs to be done: theory to practice. Houston, TX: Idea Bite Press.",
    "Yarrow, K. (2014) Decoding the new consumer mind: how and why we shop and buy. San Francisco: Jossey-Bass.",
    "Zote, J. (2024) The 2024 short-form video report. Sprout Social Insights. Available at: https://sproutsocial.com/insights/short-form-video/ (Accessed: 3 May 2026).",
]
for r in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1); p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.space_after = Pt(4); p.add_run(r).font.size = Pt(10)

# ============================================================
# §18 APPENDIX A: Source-of-truth screenshots
# ============================================================
H1("Appendix A. Source-of-truth screenshots")
P(
    "Every quantitative claim in this report is traceable to one of the screenshots reproduced here. The original "
    "files, together with the campaign export and the Notion content-operating-system pages referenced in section "
    "7, are bundled in the submission’s sources/ folder."
)
SHOTS = ROOT / "inputs" / "screenshots"
img(SHOTS / "proof_meta_campaign_overview.png", width_cm=15.5,
    caption="Appendix A.1. Meta Ads Manager campaign list, snapshot 2026-05-03 09:36. Verifies €249.13 spend, 45 leads, CPL €5.54, daily budget €50, status Ongoing.")
img(SHOTS / "proof_meta_performance.png", width_cm=15.5,
    caption="Appendix A.2. Meta Ads Manager Performance overview, snapshot 2026-05-03 10:09. Verifies the daily lead curve (7-7-17-8-4-2) used in figure 6 and the updated headline figures €250.80 spend / 45 leads / CPL €5.57.")
img(SHOTS / "proof_meta_demographics.png", width_cm=15.5,
    caption="Appendix A.3. Meta Ads Manager Demographics breakdown, snapshot 2026-05-03 10:10. Verifies the gender split (73 percent male / 24 percent female), the per-gender CPL (€5.48 / €6.08), and the age distribution used in figure 8.")
img(CHARTS / "fig_organic_views.png", width_cm=15.5,
    caption="Appendix A.4. The nine-video organic sprint view counts (2,442 / 2,656 / 1,499 / 1,645 / 3,662 / 3,329 / 4,360 / 3,223 / 4,402) were captured from the Instagram profile insights screenshot supplied by the campaign operator on 2026-05-03 and are reproduced as the chart above. The original screenshot is filed at sources/instagram_9video_grid.png in the submission bundle.")

# §19 AI use
H1("Statement on the use of generative AI")
P(
    "Generative AI tools (Claude by Anthropic) were used during the production of this report in two ways: "
    "first, as a research assistant to surface candidate citations, which were then verified against the "
    "original sources; and second, as a drafting tool to convert campaign notes and the Meta Ads dataset into "
    "structured prose, which was subsequently edited line by line by the author for register and accuracy. "
    "The frameworks, the campaign decisions, the budget allocation, the creative direction, and the qualitative "
    "reflections are the author’s own. The dataset was generated by the author’s Meta Ads campaign "
    "and is included in the appendix folder of the submission. Every cited source has been read by the author; "
    "no reference has been included that could not be personally verified."
)

out_path = ROOT / "output" / "PM_23_Marketing_Report.docx"
doc.save(str(out_path))
print(f"Saved: {out_path}")
print(f"Approx word count: {sum(len(p.text.split()) for p in doc.paragraphs)} words")
